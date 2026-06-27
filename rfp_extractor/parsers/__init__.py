from __future__ import annotations

from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from rfp_extractor.models import DocumentFormat, DocumentSection, ParsedDocument


class DocumentParser(ABC):
    """Abstract base class for document parsers"""

    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse document and return structured representation"""
        pass

    @abstractmethod
    def supported_formats(self) -> list[DocumentFormat]:
        """Return list of supported formats"""
        pass

    def _generate_document_id(self, file_path: Path) -> str:
        """Generate consistent document ID from file path"""
        import hashlib
        return hashlib.md5(str(file_path.absolute()).encode()).hexdigest()[:12]

    def _split_into_sections(self, text: str, headings: list[tuple[str, int, int]]) -> list[DocumentSection]:
        """Split text into sections based on headings"""
        sections = []
        for i, (heading, level, pos) in enumerate(headings):
            end_pos = headings[i + 1][2] if i + 1 < len(headings) else len(text)
            content = text[pos:end_pos].strip()

            # Extract section number if present
            import re
            number_match = re.match(r"^(\d+(?:\.\d+)*)\s+", heading)
            number = number_match.group(1) if number_match else None
            clean_heading = re.sub(r"^\d+(?:\.\d+)*\s+", "", heading).strip()

            sections.append(DocumentSection(
                number=number,
                heading=clean_heading,
                level=level,
                content=content,
                start_position=pos,
                end_position=end_pos
            ))
        return sections


class MarkdownParser(DocumentParser):
    """Parse Markdown files"""

    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.MARKDOWN]

    def parse(self, file_path: Path) -> ParsedDocument:
        import markdown
        from bs4 import BeautifulSoup

        content = file_path.read_text(encoding="utf-8")

        # Convert to HTML for parsing
        html = markdown.markdown(content, extensions=["fenced_code", "tables", "toc"])
        soup = BeautifulSoup(html, "html.parser")

        # Extract text with positions
        full_text = soup.get_text("\n")
        word_count = len(full_text.split())
        char_count = len(full_text)

        # Find headings with positions
        headings = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            level = int(tag.name[1])
            heading_text = tag.get_text(strip=True)
            # Find position in original text (approximate)
            pos = full_text.find(heading_text)
            if pos >= 0:
                headings.append((heading_text, level, pos))

        # Sort by position
        headings.sort(key=lambda x: x[2])

        sections = self._split_into_sections(full_text, headings)

        # If no headings found, treat whole document as one section
        if not sections:
            sections = [DocumentSection(
                heading=None,
                level=1,
                content=full_text,
                start_position=0,
                end_position=len(full_text)
            )]

        doc_id = self._generate_document_id(file_path)
        title = headings[0][0] if headings else file_path.stem

        return ParsedDocument(
            document_id=doc_id,
            format=DocumentFormat.MARKDOWN,
            source_path=str(file_path),
            title=title,
            full_text=full_text,
            sections=sections,
            word_count=word_count,
            char_count=char_count
        )


class HTMLParser(DocumentParser):
    """Parse HTML files"""

    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.HTML]

    def parse(self, file_path: Path) -> ParsedDocument:
        from bs4 import BeautifulSoup

        content = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(content, "html.parser")

        full_text = soup.get_text("\n")
        word_count = len(full_text.split())
        char_count = len(full_text)

        # Extract headings
        headings = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            level = int(tag.name[1])
            heading_text = tag.get_text(strip=True)
            pos = full_text.find(heading_text)
            if pos >= 0:
                headings.append((heading_text, level, pos))

        headings.sort(key=lambda x: x[2])
        sections = self._split_into_sections(full_text, headings)

        if not sections:
            sections = [DocumentSection(
                heading=None,
                level=1,
                content=full_text,
                start_position=0,
                end_position=len(full_text)
            )]

        title = soup.find("title")
        doc_id = self._generate_document_id(file_path)

        return ParsedDocument(
            document_id=doc_id,
            format=DocumentFormat.HTML,
            source_path=str(file_path),
            title=title.get_text(strip=True) if title else file_path.stem,
            full_text=full_text,
            sections=sections,
            word_count=word_count,
            char_count=char_count
        )


class PDFParser(DocumentParser):
    """Parse PDF files using pdfplumber"""

    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.PDF]

    def parse(self, file_path: Path) -> ParsedDocument:
        import pdfplumber

        full_text = ""
        sections = []
        page_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text:
                    full_text += f"\n--- Page {page_num} ---\n{text}"

                # Try to extract tables
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        table_text = "\n".join([" | ".join([cell or "" for cell in row]) for row in table])
                        full_text += f"\n--- Table (Page {page_num}) ---\n{table_text}"

        word_count = len(full_text.split())
        char_count = len(full_text)
        doc_id = self._generate_document_id(file_path)

        # Simple section detection for PDF (by page or major headings)
        sections = self._detect_pdf_sections(full_text)

        return ParsedDocument(
            document_id=doc_id,
            format=DocumentFormat.PDF,
            source_path=str(file_path),
            title=file_path.stem,
            full_text=full_text.strip(),
            sections=sections,
            page_count=page_count,
            word_count=word_count,
            char_count=char_count
        )

    def _detect_pdf_sections(self, text: str) -> list[DocumentSection]:
        """Detect sections in PDF text by heading patterns"""
        import re

        # Common heading patterns in RFPs
        patterns = [
            (r"^(\d+(?:\.\d+)*)\s+(.+)$", 1),  # Numbered sections
            (r"^(Section\s+\d+[:\s].+)$", 2),   # "Section 1: Title"
            (r"^(Article\s+[IVX]+[:\s].+)$", 2), # "Article I: Title"
            (r"^(Appendix\s+[A-Z][:\s].+)$", 2), # "Appendix A: Title"
            (r"^([A-Z][A-Z\s]{10,})$", 3),      # ALL CAPS headings
        ]

        headings = []
        for pattern, level in patterns:
            for match in re.finditer(pattern, text, re.MULTILINE):
                if len(match.groups()) >= 2:
                    number = match.group(1)
                    heading = match.group(2)
                else:
                    number = None
                    heading = match.group(1)
                pos = match.start()
                headings.append((heading.strip(), level, pos, number))

        headings.sort(key=lambda x: x[2])
        sections = []

        for i, (heading, level, pos, number) in enumerate(headings):
            end_pos = headings[i + 1][2] if i + 1 < len(headings) else len(text)
            content = text[pos:end_pos].strip()
            sections.append(DocumentSection(
                number=number,
                heading=heading,
                level=level,
                content=content,
                start_position=pos,
                end_position=end_pos
            ))

        if not sections:
            sections = [DocumentSection(
                heading=None,
                level=1,
                content=text,
                start_position=0,
                end_position=len(text)
            )]

        return sections


class DOCXParser(DocumentParser):
    """Parse DOCX files using python-docx"""

    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.DOCX]

    def parse(self, file_path: Path) -> ParsedDocument:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(file_path)
        full_text = ""
        headings = []
        position = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if it's a heading
            style = para.style.name.lower() if para.style else ""
            level = 1
            if "heading" in style:
                try:
                    level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
                except (ValueError, IndexError):
                    level = 1
            elif "title" in style:
                level = 1

            # Check for heading by style
            is_heading = "heading" in style or "title" in style

            pos = full_text.find(text)
            if pos == -1:
                pos = position

            if is_heading and len(text) < 200:  # Reasonable heading length
                headings.append((text, level, pos))

            full_text += text + "\n"
            position = len(full_text)

        word_count = len(full_text.split())
        char_count = len(full_text)
        doc_id = self._generate_document_id(file_path)

        headings.sort(key=lambda x: x[2])
        sections = self._split_into_sections(full_text, headings)

        if not sections:
            sections = [DocumentSection(
                heading=None,
                level=1,
                content=full_text,
                start_position=0,
                end_position=len(full_text)
            )]

        title = doc.core_properties.title or file_path.stem

        return ParsedDocument(
            document_id=doc_id,
            format=DocumentFormat.DOCX,
            source_path=str(file_path),
            title=title,
            full_text=full_text.strip(),
            sections=sections,
            word_count=word_count,
            char_count=char_count
        )


class DocumentParserFactory:
    """Factory for creating document parsers"""

    _parsers: dict[DocumentFormat, DocumentParser] = {}

    @classmethod
    def register(cls, parser: DocumentParser) -> None:
        for fmt in parser.supported_formats():
            cls._parsers[fmt] = parser

    @classmethod
    def get_parser(cls, format: DocumentFormat) -> DocumentParser | None:
        return cls._parsers.get(format)

    @classmethod
    def parse_file(cls, file_path: Path) -> ParsedDocument:
        format_map = {
            ".md": DocumentFormat.MARKDOWN,
            ".markdown": DocumentFormat.MARKDOWN,
            ".html": DocumentFormat.HTML,
            ".htm": DocumentFormat.HTML,
            ".pdf": DocumentFormat.PDF,
            ".docx": DocumentFormat.DOCX,
        }

        ext = file_path.suffix.lower()
        fmt = format_map.get(ext)

        if not fmt:
            raise ValueError(f"Unsupported file format: {ext}")

        parser = cls.get_parser(fmt)
        if not parser:
            raise ValueError(f"No parser registered for format: {fmt}")

        return parser.parse(file_path)


# Register default parsers
DocumentParserFactory.register(MarkdownParser())
DocumentParserFactory.register(HTMLParser())
DocumentParserFactory.register(PDFParser())
DocumentParserFactory.register(DOCXParser())